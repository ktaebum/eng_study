# -*- coding: utf-8 -*-

import random
import warnings

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from datetime import datetime

from ..word_io import print_file_list, read_csv


class TestGenerator(object):
    def __init__(self, path, section2=True, section3=False):
        """
        :param path: file path
        :param section2: enable section2
                        Section2 is matching word for given meaning
        :param section3: enable section3
                        Section3 is multiple choice question for given example
        """
        word_files = print_file_list(path)
        self.word_list = []
        number = input('File Numbers: ')
        number = list(map(lambda nu: int(nu), number.split(' ')))
        for n in number:
            try:
                filename = word_files[n - 1]
                self.word_list += read_csv(filename)
            except IndexError:
                print('Error: Non Existing File Number Included')

        random.shuffle(self.word_list)

        self.answer_doc = None
        self.test_doc = None
        self.question_number = 1
        self.section2 = section2
        self.section3 = section3

    def generate_tests(self, verbose=True):
        """
        section1:  Write meaning for given word
        section2:  Match word for given meaning table set
        section3:  Match word for given sentence
        """

        # Test Preface
        print('There are %d words in total' % (len(self.word_list)))
        num_questions = int(input('Set the number of questions: '))
        if not self.section2 and not self.section3:
            section1_nums = num_questions
            section2_nums = 0
            section3_nums = 0
        else:
            section1_nums = int(input('Number of questions for section 1: '))
            if self.section2:
                section2_nums = int(input('Number of questions for section 2: '))
            else:
                section2_nums = 0
            if self.section3:
                section3_nums = int(input('Number of questions for section 3: '))
            else:
                section3_nums = 0

        assert section1_nums + section2_nums + section3_nums == num_questions, 'Number of questions does not match!'

        if verbose:
            print('\n\n')
            print('===== Test Info =====')
            print('Total Questions: %d' % num_questions)
            print('Section1: %d' % section1_nums)
            print('Section2: %d' % section2_nums)
            print('Section3: %d' % section3_nums)
            print('=====================')
            print('\n\n')

        (section1, section2, section3) = self.split_sections(section1_nums, section2_nums, section3_nums)

        self.answer_doc = Document()
        self.test_doc = Document()

        font_dict = {
            'font': 'Times New Roman',
            'answer_doc_font_size': Pt(9),
            'test_doc_font_size': Pt(10)
        }

        margin_dict = {
            'top': Inches(0.5),
            'bottom': Inches(0.5)
        }

        self.set_basic_document_setting(font_dict=font_dict,
                                        margin_dict=margin_dict)
        if verbose:
            print('Making section1. . . ')

        self.make_section1(section1)

        if verbose:
            print('Making section1 finished')
            print()
            print('Making section2. . .')
        self.make_section2(section2)

        if verbose:
            print('Making section2 finished')
            print()
            print('Making section3. . .')
        self.make_section3(section3)

        if verbose:
            print('Making section3 finished')
            print()
            print('Now save into word file. . .')

        self.answer_doc.save('answer.docx')
        self.test_doc.save('test.docx')

        if verbose:
            print('Saving finished!')

        return

    def set_basic_document_setting(self, font_dict, margin_dict):
        font = font_dict.setdefault('font', 'Times New Roman')
        answer_doc_font_size = font_dict.setdefault('answer_doc_font_size', 9)
        test_doc_font_size = font_dict.setdefault('test_doc_font_size', 10)

        top_margin = margin_dict.setdefault('top', Inches(0.5))
        bottom_margin = margin_dict.setdefault('bottom', Inches(0.5))

        answer_sections = self.answer_doc.sections
        test_sections = self.test_doc.sections
        for answer_section, test_section in zip(answer_sections, test_sections):
            answer_section.top_margin = top_margin
            answer_section.bottom_margin = bottom_margin
            test_section.top_margin = top_margin
            test_section.bottom_margin = bottom_margin

        # Basic Font Setting
        answer_font = self.answer_doc.styles['Normal'].font
        test_font = self.test_doc.styles['Normal'].font
        answer_font.name = font
        answer_font.size = answer_doc_font_size
        test_font.name = font
        test_font.size = test_doc_font_size

        test_day = datetime.now()
        answer_header = self.answer_doc.add_paragraph()
        answer_header.style = self.answer_doc.styles['Heading 2']
        answer_header.add_run('Answer')
        answer_header.add_run(
            text='\t\t\t\t\t\t\t\t\t%d-%d-%d' % (test_day.year, test_day.month, test_day.day + 1))

        test_header = self.test_doc.add_paragraph()
        test_header.style = self.test_doc.styles['Heading 2']
        test_header.add_run('Test')
        test_header.add_run(
            text='\t\t\t\t\t\t\t\t\t\t%d-%d-%d' % (test_day.year, test_day.month, test_day.day + 1))

    def make_section1(self, section1):
        if len(section1) == 0:
            return

        answer = self.answer_doc.add_paragraph()
        test = self.test_doc.add_paragraph()
        answer.add_run('Answer for section 1').bold = True
        test.add_run('Section 1: Write ').bold = True
        meaning = test.add_run('meaning')
        meaning.bold = True
        meaning.italic = True
        test.add_run(' and at least one ').bold = True
        synonym = test.add_run('synonym')
        synonym.bold = True
        synonym.italic = True
        test.add_run(' of given words').bold = True
        # test.add_run('Section 1: Write meaning and at least one synonym of given words').bold = True

        for word in section1:
            self.answer_doc.add_paragraph('%d) %s / %s' % (self.question_number, word.e_mean, word.k_mean))
            self.test_doc.add_paragraph('%d) %s' % (self.question_number, word.word))
            self.question_number += 1

        return

    def make_section2(self, section2):
        if len(section2) == 0:
            return
        self.test_doc.add_page_break()
        answer = self.answer_doc.add_paragraph()
        test = self.test_doc.add_paragraph()
        answer.add_run('Answer for section 2').bold = True
        test.add_run('Section 2: Match word for given meaning').bold = True

        self.test_doc.add_paragraph('============================================================================')
        word_index = ord('a')
        p = None
        for i, word in enumerate(section2):
            word.index = chr(word_index)
            if i % 5 == 0:
                p = self.test_doc.add_paragraph()
            p.add_run('(%s) %s \t' % (word.index, word))
            word_index += 1

        random.shuffle(section2)

        self.test_doc.add_paragraph('============================================================================')
        for word in section2:
            self.test_doc.add_paragraph('%d) %s / %s' % (self.question_number, word.e_mean, word.k_mean))
            self.answer_doc.add_paragraph('%d) (%s)-%s' % (self.question_number, word.index, word.word))
            self.question_number += 1

        return

    def make_section3(self, section3):
        if len(section3) == 0:
            return

        self.test_doc.add_page_break()
        answer = self.answer_doc.add_paragraph()
        test = self.test_doc.add_paragraph()
        answer.add_run('Answer for section 3').bold = True
        test.add_run('Section 3: Select ').bold = True
        appropriate = test.add_run('the most appropriate')
        appropriate.bold = True
        appropriate.italic = True
        test.add_run(' word for given sentence from following choices').bold = True
        choices = ['a', 'b', 'c', 'd']
        for word in section3:
            self.test_doc.add_paragraph('%d) %s' % (self.question_number, word.sentence))
            answer_choice = random.choice(choices)
            choice = self.test_doc.add_paragraph('')
            others = random.sample(list(filter(lambda w: w.choice_appear < 4 and w != word, section3)),
                                   k=len(choices) - 1)
            for c in choices:
                if c == answer_choice:
                    choice.add_run('(%s) %s\t' % (c, word.word))
                    self.answer_doc.add_paragraph('%d) (%s)-%s' % (self.question_number, c, word.word))
                    word.choice_appear += 1
                else:
                    other = others.pop()
                    choice.add_run('(%s) %s\t' % (c, other.word))
                    other.choice_appear += 1
            self.question_number += 1

        return

    def split_sections(self, section1_nums, section2_nums, section3_nums):
        """
        split each sections
        :param section1_nums: # of questions for section1
        :param section2_nums: # of questions for section2
        :param section3_nums: # of questions for section3
        :return: section1, section2, section3 words in list
        """
        word_list = self.word_list.copy()
        if self.section3:
            filtered_section3 = list(filter(lambda w: w.sentence != '', word_list))
            try:
                section3 = random.sample(filtered_section3,
                                         k=section3_nums)
            except ValueError:
                # Catch when section3_nums is greater than length of filtered_section3
                warnings.warn(
                    'There are only %d words which have example! Make section 3 with these %d words, not %d words'
                    % (len(filtered_section3), section3_nums, len(filtered_section3)),
                    RuntimeWarning
                )
                section3 = filtered_section3
                section1_nums += section3_nums - len(section3)
            for word in section3:
                # Remove section3 words from original word list
                word_list.remove(word)
                # word_list = list(filter(lambda w: not w.equal(word), word_list))
            # Replace word to ______
            section3 = list(map(lambda w: w.remove_word_from_sentence(), section3))
        else:
            section3 = []

        section1 = word_list[:section1_nums]
        section2 = word_list[section1_nums:section1_nums + section2_nums]
        return section1, section2, section3
